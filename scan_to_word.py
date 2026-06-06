#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Сканирует текущую директорию, собирает файлы .go, .sql, .env
и создаёт Word-документ с их содержимым.
Зависимость: pip install python-docx
"""

import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("Ошибка: не установлена библиотека python-docx")
    print("Установите её командой: pip install python-docx")
    sys.exit(1)

def collect_files(root_dir):
    """Рекурсивно собирает пути к файлам .go, .sql, .env"""
    extensions = {'.go', '.sql', '.env'}
    found_files = []

    for dirpath, _, filenames in os.walk(root_dir):
        for filename in filenames:
            # Проверяем расширение (включая .env, .env.local и т.д.)
            if any(filename.endswith(ext) for ext in extensions):
                full_path = os.path.join(dirpath, filename)
                # Относительный путь от корня сканирования
                rel_path = os.path.relpath(full_path, start=root_dir)
                found_files.append((rel_path, full_path))

    return found_files

def add_code_paragraph(doc, text):
    """Добавляет параграф с моноширинным шрифтом и сохраняет отступы"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # Сохраняем пробелы и переносы строк (Word их обычно сохраняет)
    # Если нужно строгое форматирование кода, лучше вставлять как preformatted,
    # но python-docx не поддерживает прямые <pre>. Однако обычный текст с пробелами
    # и табуляцией отображается корректно, если шрифт моноширинный.
    return p

def create_word_document(files_data, output_path="directory_content.docx"):
    """Создаёт Word-документ со структурой: путь -> содержимое файла"""
    doc = Document()
    doc.add_heading("Содержимое директории", level=1)
    doc.add_paragraph(f"Сканирование корня: {os.getcwd()}")
    doc.add_paragraph(f"Найдено файлов: {len(files_data)}")

    for rel_path, full_path in files_data:
        # Заголовок с путём
        doc.add_heading(rel_path, level=2)

        # Чтение содержимого файла
        try:
            with open(full_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            # Попробуем другую кодировку (например, cp1251)
            try:
                with open(full_path, 'r', encoding='cp1251') as f:
                    content = f.read()
            except Exception as e:
                content = f"// Не удалось прочитать файл: {e}"
        except Exception as e:
            content = f"// Ошибка доступа: {e}"

        # Вставляем содержимое как моноширинный текст
        add_code_paragraph(doc, content)

        # Разделитель между файлами (необязательно)

    # Сохраняем документ
    doc.save(output_path)
    print(f"Документ успешно создан: {output_path}")

def main():
    # Директория, в которой запущен скрипт
    root = os.getcwd()
    print(f"Сканируем директорию: {root}")

    files = collect_files(root)
    if not files:
        print("Не найдено файлов с расширениями .go, .sql или .env")
        # Всё равно создадим документ с сообщением
        doc = Document()
        doc.add_heading("Содержимое директории", level=1)
        doc.add_paragraph(f"Сканирование: {root}")
        doc.add_paragraph("Нет подходящих файлов (.go, .sql, .env).")
        doc.save("no_files_found.docx")
        print("Создан пустой отчёт: no_files_found.docx")
        return

    print(f"Найдено файлов: {len(files)}")
    create_word_document(files)

if __name__ == "__main__":
    main()